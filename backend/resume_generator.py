"""
Resume Generator - PDF and DOCX output using ReportLab and python-docx
"""
import os, tempfile, re
from datetime import datetime

# ── PDF Generation ─────────────────────────────────────────────────────────────
def generate_pdf_resume(data: dict) -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, Table, TableStyle, KeepTogether)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

    p = data.get("personal", {})
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", dir="/tmp")
    tmp.close()

    doc = SimpleDocTemplate(
        tmp.name,
        pagesize=letter,
        leftMargin=0.65*inch, rightMargin=0.65*inch,
        topMargin=0.6*inch, bottomMargin=0.6*inch
    )

    NAVY  = colors.HexColor("#1a2e4a")
    BLUE  = colors.HexColor("#2563eb")
    GRAY  = colors.HexColor("#64748b")
    LGRAY = colors.HexColor("#f1f5f9")
    BLACK = colors.HexColor("#0f172a")

    styles = getSampleStyleSheet()

    def style(name, **kw):
        return ParagraphStyle(name, **kw)

    S_NAME    = style("Name",    fontName="Helvetica-Bold", fontSize=22, textColor=NAVY, leading=26, alignment=TA_CENTER)
    S_CONTACT = style("Contact", fontName="Helvetica",      fontSize=9,  textColor=GRAY, leading=12, alignment=TA_CENTER)
    S_SEC     = style("Section", fontName="Helvetica-Bold", fontSize=11, textColor=BLUE, leading=14, spaceBefore=8, spaceAfter=2)
    S_TITLE   = style("Title",   fontName="Helvetica-Bold", fontSize=10, textColor=BLACK, leading=13)
    S_SUB     = style("Sub",     fontName="Helvetica",      fontSize=9,  textColor=GRAY, leading=12)
    S_BODY    = style("Body",    fontName="Helvetica",      fontSize=9,  textColor=BLACK, leading=13, leftIndent=8)
    S_BULLET  = style("Bullet",  fontName="Helvetica",      fontSize=9,  textColor=BLACK, leading=13, leftIndent=16, firstLineIndent=-8)
    S_SKILLS  = style("Skills",  fontName="Helvetica",      fontSize=9,  textColor=BLACK, leading=13)

    story = []

    # ── Header ──
    story.append(Paragraph(p.get("name", "Your Name"), S_NAME))
    contacts = []
    for field in ["email","phone","location","linkedin","github","portfolio"]:
        val = p.get(field, "")
        if val:
            contacts.append(val)
    story.append(Paragraph("  •  ".join(contacts), S_CONTACT))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=2, color=NAVY))
    story.append(Spacer(1, 4))

    def section(title):
        story.append(Paragraph(title.upper(), S_SEC))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1")))
        story.append(Spacer(1, 3))

    def bullet_item(text):
        story.append(Paragraph(f"• {text}", S_BULLET))

    # ── Summary ──
    summary = data.get("summary", "")
    if summary:
        section("Professional Summary")
        story.append(Paragraph(summary, S_BODY))
        story.append(Spacer(1, 4))

    # ── Experience ──
    experiences = data.get("experience", [])
    if experiences:
        section("Work Experience")
        for exp in experiences:
            title_line = f"<b>{exp.get('title','')}</b>"
            company_line = f"{exp.get('company','')}  |  {exp.get('location','')}  |  {exp.get('duration','')}"
            story.append(Paragraph(title_line, S_TITLE))
            story.append(Paragraph(company_line, S_SUB))
            for bp in exp.get("bullets", []):
                if bp.strip():
                    bullet_item(bp)
            story.append(Spacer(1, 5))

    # ── Education ──
    educations = data.get("education", [])
    if educations:
        section("Education")
        for edu in educations:
            row1 = f"<b>{edu.get('degree','')}</b> — {edu.get('institution','')}"
            row2 = f"{edu.get('year','')}  |  GPA: {edu.get('gpa','')}" if edu.get('gpa') else edu.get('year','')
            story.append(Paragraph(row1, S_TITLE))
            story.append(Paragraph(row2, S_SUB))
            story.append(Spacer(1, 4))

    # ── Projects ──
    projects = data.get("projects", [])
    if projects:
        section("Projects")
        for proj in projects:
            name = proj.get("name", "")
            tech = proj.get("tech", "")
            desc = proj.get("description", "")
            link = proj.get("link", "")
            title_str = f"<b>{name}</b>"
            if tech:
                title_str += f"  <font color='#64748b' size='8'>| {tech}</font>"
            story.append(Paragraph(title_str, S_TITLE))
            if desc:
                story.append(Paragraph(desc, S_BODY))
            if link:
                story.append(Paragraph(f"<font color='#2563eb'>{link}</font>", S_SUB))
            story.append(Spacer(1, 5))

    # ── Skills ──
    skills = data.get("skills", {})
    if skills:
        section("Technical Skills")
        for category, skill_list in skills.items():
            if skill_list:
                items = skill_list if isinstance(skill_list, list) else [skill_list]
                story.append(Paragraph(f"<b>{category}:</b>  {', '.join(items)}", S_SKILLS))
                story.append(Spacer(1, 2))

    # ── Certifications ──
    certs = data.get("certifications", [])
    if certs:
        section("Certifications")
        for c in certs:
            story.append(Paragraph(f"• {c}", S_BULLET))

    # ── Achievements ──
    achievements = data.get("achievements", [])
    if achievements:
        section("Achievements & Awards")
        for a in achievements:
            story.append(Paragraph(f"• {a}", S_BULLET))

    doc.build(story)
    return tmp.name


# ── DOCX Generation ────────────────────────────────────────────────────────────
def generate_docx_resume(data: dict) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    NAVY  = RGBColor(0x1a, 0x2e, 0x4a)
    BLUE  = RGBColor(0x25, 0x63, 0xeb)
    GRAY  = RGBColor(0x64, 0x74, 0x8b)

    p = data.get("personal", {})
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    def set_font(run, size=10, bold=False, color=None, italic=False):
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color

    def add_name():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(p.get("name", "Your Name"))
        set_font(run, size=22, bold=True, color=NAVY)

    def add_contact():
        parts = [p.get(f, "") for f in ["email","phone","location","linkedin","github"] if p.get(f)]
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("  •  ".join(parts))
        set_font(run, size=9, color=GRAY)

    def add_section_header(title):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(2)
        run = para.add_run(title.upper())
        set_font(run, size=11, bold=True, color=BLUE)
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2563eb')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text):
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.left_indent = Inches(0.2)
        run = para.add_run(text)
        set_font(run, size=9)

    add_name()
    add_contact()

    summary = data.get("summary", "")
    if summary:
        add_section_header("Professional Summary")
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        run = p2.add_run(summary)
        set_font(run, size=9)

    for exp in data.get("experience", []):
        add_section_header("Work Experience") if exp == data["experience"][0] else None
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(exp.get("title", ""))
        set_font(r, size=10, bold=True)
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(2)
        r2 = p3.add_run(f"{exp.get('company','')}  |  {exp.get('location','')}  |  {exp.get('duration','')}")
        set_font(r2, size=9, color=GRAY)
        for bp in exp.get("bullets", []):
            if bp.strip():
                add_bullet(bp)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for edu in data.get("education", []):
        add_section_header("Education") if edu == data["education"][0] else None
        p2 = doc.add_paragraph()
        r = p2.add_run(f"{edu.get('degree','')} — {edu.get('institution','')}")
        set_font(r, size=10, bold=True)
        p3 = doc.add_paragraph()
        r2 = p3.add_run(f"{edu.get('year','')}  |  GPA: {edu.get('gpa','')}" if edu.get('gpa') else edu.get('year',''))
        set_font(r2, size=9, color=GRAY)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for proj in data.get("projects", []):
        add_section_header("Projects") if proj == data["projects"][0] else None
        p2 = doc.add_paragraph()
        r = p2.add_run(proj.get("name",""))
        set_font(r, size=10, bold=True)
        if proj.get("tech"):
            r2 = p2.add_run(f"  |  {proj['tech']}")
            set_font(r2, size=9, color=GRAY)
        if proj.get("description"):
            p3 = doc.add_paragraph()
            r3 = p3.add_run(proj["description"])
            set_font(r3, size=9)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    skills = data.get("skills", {})
    if skills:
        add_section_header("Technical Skills")
        for category, items in skills.items():
            if items:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(2)
                r1 = p2.add_run(f"{category}: ")
                set_font(r1, size=9, bold=True)
                vals = items if isinstance(items, list) else [items]
                r2 = p2.add_run(", ".join(vals))
                set_font(r2, size=9)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir="/tmp")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name
